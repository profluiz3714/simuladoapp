import os
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.contrib.auth import login, authenticate
from django.contrib import messages
from docx import Document
import pandas as pd
from .models import Questao
from .forms import QuestaoForm, RegistroForm
from bs4 import BeautifulSoup
from django.conf import settings
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
import random
import tempfile
import zipfile

def index(request):
    questoes = Questao.objects.all()
    return render(request, 'questoes/index.html', {'questoes': questoes})

def adicionar_questao(request):
    if request.method == 'POST':
        form = QuestaoForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Questão adicionada com sucesso.')
            return redirect('adicionar_questao')
        else:
            messages.error(request, 'Erro ao adicionar a questão. Verifique os dados informados.')
    else:
        form = QuestaoForm()
    return render(request, 'questoes/adicionar_questao.html', {'form': form})

def editar_questao(request, id):
    questao = get_object_or_404(Questao, id=id)
    if request.method == 'POST':
        form = QuestaoForm(request.POST, request.FILES, instance=questao)
        if form.is_valid():
            form.save()
            return redirect('index')
    else:
        form = QuestaoForm(instance=questao)
    return render(request, 'questoes/editar_questao.html', {'form': form})

def excluir_questao(request, id):
    questao = get_object_or_404(Questao, id=id)
    if request.method == 'POST':
        questao.delete()
        return redirect('index')
    return render(request, 'questoes/excluir_questao.html', {'questao': questao})

def gerar_documento_simulado(tipo, questoes, correlacoes, tipo_col):
    document = Document()

    # Configurar margens estreitas
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Configurar layout em duas colunas
    sectPr = sections[0]._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:sep'), '1')

    # Adicionar o título
    document.add_heading(f'Simulado {tipo}', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, questao in enumerate(questoes, 1):
        p = document.add_paragraph(f'Questão {i}')
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(11)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        soup = BeautifulSoup(questao.enunciado, 'html.parser')
        
        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                p = document.add_paragraph()
                p.add_run(element.get_text())
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(11)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif element.name == 'img':
                image_path = element['src']
                if image_path.startswith('/media/'):
                    image_path = os.path.join(settings.MEDIA_ROOT, image_path[len(settings.MEDIA_URL):])
                document.add_picture(image_path, width=Inches(3.0))

        correlacoes.loc[questao.id, tipo_col] = i  # Usar o ID da questão para manter a correlação
        if tipo_col == 'Padrão':
            correlacoes.loc[questao.id, 'Disciplina'] = questao.disciplina
            correlacoes.loc[questao.id, 'Gabarito'] = questao.gabarito

    return document

def gerar_simulado(request):
    tipos = ['Padrão', 'Azul', 'Amarelo', 'Cinza', 'Rosa']
    questoes = list(Questao.objects.all())
    correlacoes = pd.DataFrame(columns=['Disciplina', 'Gabarito'] + tipos)

    with tempfile.TemporaryDirectory() as tempdir:
        zip_filename = os.path.join(tempdir, 'simulados.zip')
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for tipo in tipos:
                questoes_em_ordem = random.sample(questoes, len(questoes))  # Embaralhar as questões
                document = gerar_documento_simulado(tipo, questoes_em_ordem, correlacoes, tipo)
                temp_file_path = os.path.join(tempdir, f'simulado_{tipo}.docx')
                document.save(temp_file_path)
                zipf.write(temp_file_path, f'simulado_{tipo}.docx')

            # Gerar arquivo Excel com correlações
            correlacoes.reset_index(inplace=True)
            correlacoes.rename(columns={'index': 'Questão_ID'}, inplace=True)
            correlacoes.set_index('Questão_ID', inplace=True)
            correlacoes.to_excel(os.path.join(tempdir, 'correlacoes.xlsx'))

            zipf.write(os.path.join(tempdir, 'correlacoes.xlsx'), 'correlacoes.xlsx')

        # Enviar o arquivo ZIP como resposta
        with open(zip_filename, 'rb') as zip_file:
            response = HttpResponse(zip_file.read(), content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename=simulados.zip'
            return response
    pass

def gerar_planilha(request):
    questoes = Questao.objects.all()
    data = []

    for questao in questoes:
        data.append({
            'Enunciado': questao.enunciado,
            'Disciplina': questao.disciplina,
            'Gabarito': questao.gabarito
        })

    df = pd.DataFrame(data)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=questoes.xlsx'
    df.to_excel(response, index=False)
    return response

def recomecar(request):
    Questao.objects.all().delete()
    return redirect('listar_questoes')

def registro(request):
    if request.method == 'POST':
        form = RegistroForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password1')
            user = authenticate(username=username, password=password)
            login(request, user)
            messages.success(request, 'Usuário cadastrado com sucesso!')
            return redirect('index')
    else:
        form = RegistroForm()
    return render(request, 'questoes/registro.html', {'form': form})

def listar_questoes(request):
    questoes = Questao.objects.all()
    return render(request, 'questoes/listar_questoes.html', {'questoes': questoes})
